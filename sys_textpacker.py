import os
from docx import Document

num=0
#data=""
index=0

def docx_pack(line,index):
    
    doc.add_paragraph(data)
    doc.save(str(index)+".docx")


file=open("sys_text.txt",'r')
all_data=file.readlines()
doc = Document()
for line in all_data:
    line=line.replace('◙','')
    doc.add_paragraph(line[0:-1])
    num=num+1
    if num==4999:
        index=index+1
        doc.save(str(index)+".docx")
        num=0
        #docx_pack(data,index)
        doc = Document()
index=index+1
doc.save(str(index)+".docx")
